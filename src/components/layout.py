import streamlit as st
from docx import Document
from bs4 import BeautifulSoup
import requests
import docx
from io import BytesIO
import streamlit.components.v1 as components
from streamlit_lexical import streamlit_lexical

def page_header():
    with st.container():
        st.title("CARAble")
        st.markdown(
            "<p style='margin-top: -10px; margin-bottom: 0.5rem;'>Enable your content with CARA — your Content Authoring & Review Assistant for better content governance</p>",
            unsafe_allow_html=True,
        )
    st.divider()

def add_sidebar_copyright():
    """Add copyright information to the left sidebar at the bottom"""
    with st.sidebar:
        st.markdown(
            """
            <div style='text-align: center; padding: 1rem 0; color: #1e40af; font-size: 0.85rem;'>
                <p style='margin: 0; font-weight: 500;'>© 2025 <strong>Afiqah Rashid</strong></p>
                <p style='margin: 0.25rem 0 0 0; font-size: 0.75rem; color: #6b7280;'>CARAble</p>
            </div>
            """,
            unsafe_allow_html=True
        )

def input_section():
    with st.container():
        col1, col2 = st.columns([1, 1], gap="large")

        with col1:
            st.markdown("### Start your content review")
            
            # Add content type selection
            input_type = st.radio(
                "Choose how you'd like to provide your draft",
                ("Paste your draft directly", "Website URL", "Upload a Word document"),
                horizontal=False
            )
            
            content = ""
            
            if input_type == "Paste your draft directly":
                st.markdown("**Paste your content below:**")
                content = streamlit_lexical(
                    value="",
                    placeholder="Paste your draft content here...", 
                    height=250,
                    debounce=500,
                    key='content_editor', 
                    on_change=None
                )
                
                # Add character count and tips
                if content:
                    char_count = len(content)
                    st.caption(f"{char_count:,} characters")
                    if char_count < 100:
                        st.warning("Consider adding more content for better analysis")
                    elif char_count > 15000:
                        st.info("Content will be truncated to 15,000 characters for analysis")
                    elif char_count > 10000:
                        st.info("Large content detected - analysis may take longer")
                
            elif input_type == "Website URL":
                st.markdown("**Enter a gov.sg website URL:**")
                url = st.text_input(
                    "Website URL", 
                    placeholder="https://www.gov.sg/...",
                    help="Only gov.sg websites are supported for content extraction"
                )
                if url and "gov.sg" not in url.lower():
                    st.error("Only gov.sg URLs are allowed.")
                elif url:
                    with st.spinner("Extracting content from website..."):
                        content = fetch_webpage_text(url) or ""
                        if content:
                            st.success("Content extracted successfully!")
                        else:
                            st.error("Could not extract content from this URL")
                            
            elif input_type == "Upload a Word document":
                st.markdown("**Upload a Word document (.docx):**")
                uploaded_file = st.file_uploader(
                    "Choose a file", 
                    type=["docx"],
                    help="Upload a .docx file to analyse its content"
                )
                if uploaded_file:
                    with st.spinner("Reading document..."):
                        content = read_docx(uploaded_file)
                        if content:
                            st.success(f"Document loaded successfully! ({len(content)} characters)")
                        else:
                            st.error("Could not read the document content")
            
            # Submit button
            run = st.button(
                "Enable your content with CARAble", 
                use_container_width=True,
                type="primary",
                help="Click to analyse your content for governance, accessibility, and SEO"
            )

        with col2:
            disclaimer_section()
            faq_section()

    return content, run

def faq_section():
    st.markdown("### FAQs")
    with st.expander("Who is CARAble?"):
        st.markdown("""
CARAble is your helpful content assistant designed to support public officers in creating clear, citizen-friendly web pages — fast and confidently.

With CARAble, you can:
- Get smart recommendations for logical page structure and headings
- Rewrite content in a clear, professional tone suited for your audience
- Ensure your content meets WCAG accessibility standards
- Optimise your pages for SEO using trusted best practices
- Receive a detailed Governance Report Card  
- Easily compare your original and optimised content side-by-side
""")
    with st.expander("How do I get started and use CARAble?"):
        st.markdown("""
1. Submit your draft by pasting text, entering a gov.sg URL, or uploading a Word document.  
2. Click **Enable your content with CARAble** to start the review. 
3. View your content score card and detailed improvement suggestions.
4. Compare your original and optimised drafts to see the changes clearly.
5. Download your newly improved suggested copy.
""")
    with st.expander("Is CARAble still being improved?"):
        st.markdown("""
Yes, CARAble is a work in progress, and we’re committed to making it better over time to help you create clearer, more consistent, and compliant content.
""")
        
def disclaimer_section():
    st.markdown("### Disclaimer")
    with st.expander("Important Notice"):
        st.markdown("""This web application is a prototype developed for educational purposes only. The information provided here is NOT intended for real-world usage and should not be relied upon for making any decisions, especially those related to financial, legal, or healthcare matters.

Furthermore, please be aware that the LLM may generate inaccurate or incorrect information. You assume full responsibility for how you use any generated output.

Always consult with qualified professionals for accurate and personalised advice.
     """)

def fetch_webpage_text(url: str) -> str:
    headers = {"User-Agent": "Mozilla/5.0"}
    res = requests.get(url, headers=headers, timeout=10)
    res.raise_for_status()
    soup = BeautifulSoup(res.text, "html.parser")
    main_content = soup.find("main") or soup.find("body")
    if not main_content:
        return ""

    # Remove unwanted tags
    for tag in main_content(["script", "nav", "footer", "header", "noscript", "form", "img", "button", "input", "select"]):
        tag.decompose()

    # Remove social media bars (common class/id keywords)
    social_selectors = [
        '[class*="social"]',
        '[class*="share"]',
        '[class*="follow"]',
        '[class*="fb"]',
        '[class*="twitter"]',
        '[class*="instagram"]',
        '[class*="linkedin"]',
        '[class*="pinterest"]',
        '[id*="social"]',
        '[id*="share"]',
        '[id*="follow"]',
    ]

    for selector in social_selectors:
        for elem in main_content.select(selector):
            elem.decompose()
    return str(main_content)


def read_docx(uploaded_file) -> str:
    doc = Document(uploaded_file)
    full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return "\n".join(full_text)


def display_content_score_card(result: dict):
    st.markdown("### Content Governance Score Card")
    
    scores = result.get("governance_report", {})
    structure_score = scores.get("structure_score", "N/A")
    tone_score = scores.get("tone_score", "N/A")
    accessibility_score = scores.get("accessibility_score", "N/A")
    seo_score = scores.get("seo_score", "N/A")
    
    # Individual category scores with enhanced display
    cols = st.columns(4)
    category_scores = [
        ("Structure", structure_score, "Content organisation and hierarchy"),
        ("Tone", tone_score, "Writing style and tone compliance with content playbook"),
        ("Accessibility", accessibility_score, "WCAG 2.1 AA accessibility compliance"),
        ("SEO", seo_score, "SEO checks against Google Lighthouse audits")
    ]
    
    for i, (category, score, short_desc) in enumerate(category_scores):
        with cols[i]:
            # Extract numeric score for color coding
            if isinstance(score, str) and "/" in score:
                score_val = float(score.split("/")[0])
            else:
                score_val = 0
            
            # Enhanced visual indicators
            if score_val >= 8.0:
                delta_color = "normal"
                status_text = "Excellent"
            elif score_val >= 6.0:
                delta_color = "normal"
                status_text = "Good"
            else:
                delta_color = "inverse"
                status_text = "Needs Work"
            
            # Display metric directly without accordion
            st.metric(
                category, 
                f"{score}", 
                help=short_desc,
                delta_color=delta_color
            )
            st.caption(f"**Status:** {status_text}")
    
    
    # Suggestions with enhanced display
    st.markdown("### Improvement Suggestions")
    tabs = st.tabs(["Structure", "Tone", "Accessibility", "SEO"])
    
    suggestions_map = {
        "Structure": result.get("structure_fixes", result.get("recommended_structure", [])),
        "Tone": result.get("tone_fixes", []),
        "Accessibility": result.get("accessibility_fixes", []),
        "SEO": result.get("seo_fixes", [])
    }
    
    for tab, tab_label in zip(tabs, ["Structure", "Tone", "Accessibility", "SEO"]):
        with tab:
            items = suggestions_map.get(tab_label, [])
            if items and len(items) > 0:
                for i, item in enumerate(items, 1):
                    st.markdown(f"**{i}.** {item}")
            else:
                # Get the corresponding score to show why no suggestions
                score_key = f"{tab_label.lower()}_score"
                score = scores.get(score_key, "N/A")
                if isinstance(score, str) and "/" in score:
                    score_val = float(score.split("/")[0])
                    if score_val >= 9.0:
                        st.success(f"Excellent {tab_label} score ({score}) - No improvements needed!")
                    elif score_val >= 7.0:
                        st.info(f"Good {tab_label} score ({score}) - Content meets standards!")
                    else:
                        st.warning(f"Low {tab_label} score ({score}) - Analysis may have missed some issues")
                else:
                    st.info(f"No specific suggestions for {tab_label} - content meets standards!")


def get_embedded_html_with_style(html_content: str) -> str:
    embedded_css = """
    <style>
    /* CSS Custom Properties for consistent theming */
    :root {
        --primary-color: #7c3aed; /* Violet 600 - modern violet */
        --primary-hover: #6d28d9; /* Violet 700 */
        --primary-light: #ede9fe; /* Violet 100 */
        --secondary-color: #059669; /* Green 600 - for success states */
        --accent-color: #dc2626; /* Red 600 - for errors and highlights */
        --warning-color: #d97706; /* Orange 600 - for warnings */
        --text-primary: #1f2937; /* Gray 800 - high contrast text */
        --text-secondary: #4b5563; /* Gray 600 - secondary text */
        --text-muted: #6b7280; /* Gray 500 - muted text */
        --background-primary: #ffffff; /* White background */
        --background-secondary: #f9fafb; /* Gray 50 - subtle background */
        --background-accent: #f3f4f6; /* Gray 100 - accent background */
        --border-color: #e5e7eb; /* Gray 200 - subtle borders */
        --shadow-sm: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
        --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }

    .content-container {
        font-family: 'Source Sans Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen,
                     Ubuntu, Cantarell, 'Open Sans', 'Helvetica Neue', sans-serif;
        color: var(--text-primary);
        padding: 1.5rem;
        line-height: 1.7;
        max-width: 100%;
        background-color: var(--background-primary);
    }

    .content-container h1 {
        color: var(--text-primary);
        font-weight: 700;
        font-size: 2.2rem;
        margin-top: 0;
        margin-bottom: 1.5rem;
        line-height: 1.2;
        letter-spacing: -0.025em;
    }

    .content-container h2 {
        color: var(--primary-color);
        font-weight: 600;
        font-size: 1.8rem;
        margin-top: 2.5rem;
        margin-bottom: 1rem;
        line-height: 1.3;
        letter-spacing: -0.025em;
        border-bottom: 3px solid var(--primary-color);
        padding-bottom: 0.5rem;
    }

    .content-container h3 {
        color: var(--text-primary);
        font-weight: 600;
        font-size: 1.5rem;
        margin-top: 2rem;
        margin-bottom: 0.75rem;
        line-height: 1.4;
    }

    .content-container h4 {
        color: var(--text-secondary);
        font-weight: 600;
        font-size: 1.25rem;
        margin-top: 1.5rem;
        margin-bottom: 0.5rem;
        line-height: 1.4;
    }

    .content-container h5 {
        color: var(--text-secondary);
        font-weight: 600;
        font-size: 1.1rem;
        margin-top: 1.25rem;
        margin-bottom: 0.5rem;
        line-height: 1.4;
    }

    .content-container h6 {
        color: var(--text-muted);
        font-weight: 600;
        font-size: 1rem;
        margin-top: 1rem;
        margin-bottom: 0.5rem;
        line-height: 1.4;
    }

    .content-container p {
        font-size: 1.125rem;
        line-height: 1.7;
        margin-bottom: 1.25rem;
        text-align: left;
        color: var(--text-primary);
    }

    .content-container ul, .content-container ol {
        margin-bottom: 1.25rem;
        padding-left: 1.5rem;
    }

    .content-container li {
        font-size: 1.125rem;
        line-height: 1.6;
        margin-bottom: 0.5rem;
        color: var(--text-primary);
    }

    .content-container a {
        color: var(--primary-color);
        text-decoration: underline;
        text-decoration-thickness: 2px;
        text-underline-offset: 2px;
        font-weight: 500;
        transition: all 0.2s ease;
    }

    .content-container a:hover {
        color: var(--primary-hover);
        text-decoration-thickness: 3px;
    }

    .content-container a:focus {
        outline: 3px solid var(--primary-light);
        outline-offset: 2px;
        border-radius: 4px;
        text-decoration: none;
    }

    .content-container table {
        width: 100%;
        border-collapse: collapse;
        margin: 1.5rem 0;
        font-size: 1rem;
        box-shadow: var(--shadow-sm);
        border-radius: 8px;
        overflow: hidden;
    }

    .content-container th, .content-container td {
        border: 1px solid var(--border-color);
        padding: 1rem;
        text-align: left;
        vertical-align: top;
    }

    .content-container th {
        background-color: var(--background-accent);
        font-weight: 600;
        color: var(--text-primary);
        font-size: 1.1rem;
    }

    .content-container td {
        background-color: var(--background-primary);
    }

    .content-container blockquote {
        border-left: 4px solid var(--primary-color);
        margin: 1.5rem 0;
        padding: 1.25rem 1.5rem;
        background-color: var(--background-secondary);
        font-style: italic;
        border-radius: 0 8px 8px 0;
    }

    .content-container strong, .content-container b {
        font-weight: 600;
        color: var(--text-primary);
    }

    .content-container em, .content-container i {
        font-style: italic;
        color: var(--text-secondary);
    }

    .content-container code {
        background-color: var(--background-accent);
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
        font-size: 0.9rem;
        color: var(--text-primary);
        border: 1px solid var(--border-color);
    }

    .content-container pre {
        background-color: var(--background-accent);
        padding: 1rem;
        border-radius: 8px;
        overflow-x: auto;
        border: 1px solid var(--border-color);
        margin: 1rem 0;
    }

    .content-container pre code {
        background: none;
        padding: 0;
        border: none;
        font-size: 0.9rem;
    }

    /* Accessibility improvements */
    @media (prefers-reduced-motion: reduce) {
        .content-container * {
            animation-duration: 0.01ms !important;
            animation-iteration-count: 1 !important;
            transition-duration: 0.01ms !important;
        }
    }

    /* High contrast mode support */
    @media (prefers-contrast: high) {
        .content-container {
            color: #000000;
            background-color: #ffffff;
        }
        
        .content-container h1, 
        .content-container h2, 
        .content-container h3, 
        .content-container h4,
        .content-container h5,
        .content-container h6 {
            color: #000000;
        }
        
        .content-container a {
            color: #0000ff;
        }
        
        .content-container th {
            background-color: #f0f0f0;
            color: #000000;
        }
    }

    /* Print styles */
    @media print {
        .content-container {
            color: #000000;
            background-color: #ffffff;
        }
        
        .content-container a {
            color: #0000ff;
            text-decoration: underline;
        }
    }
    </style>

    """
    return f"""
    {embedded_css}
    <div class="content-container">
        {html_content}
    </div>
    """

def display_formatted_content(content: str, is_original: bool = True):
    """Display content with proper formatting instead of markdown preview"""
    if not content.strip():
        st.info("No content to display")
        return
    
    # Parse HTML content
    soup = BeautifulSoup(content, 'html.parser')
    
    # Create a container with proper styling
    container_class = "original-content" if is_original else "improved-content"

    # Enhanced CSS for better content display with improved layout
    css_styles = f"""
    <style>
    .{container_class} {{
        font-family: 'Source Sans Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        line-height: 1.7;
        color: #1f2937;
        padding: 1.5rem;
        background: {'#f8fafc' if is_original else '#f0f9ff'};
        border-radius: 12px;
        border: 1px solid {'#e5e7eb' if is_original else '#7c3aed'};
        max-height: none;
        overflow: visible;
        margin: 0.5rem 0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
        width: 100% !important;
        min-width: 100%;
        max-width: 100%;
        word-wrap: break-word;
        overflow-wrap: break-word;
    }}
    
    .{container_class} h1 {{
        color: #1f2937;
        font-size: 2.25rem;
        font-weight: 700;
        margin: 0 0 1.5rem 0;
        line-height: 1.2;
        border-bottom: 3px solid {'#6b7280' if is_original else '#7c3aed'};
        padding-bottom: 0.75rem;
    }}

    .{container_class} h2 {{
        color: {'#374151' if is_original else '#7c3aed'};
        font-size: 1.75rem;
        font-weight: 600;
        margin: 2rem 0 1rem 0;
        line-height: 1.3;
        border-bottom: 2px solid {'#d1d5db' if is_original else '#a78bfa'};
        padding-bottom: 0.5rem;
    }}

    .{container_class} h3 {{
        color: #1f2937;
        font-size: 1.4rem;
        font-weight: 600;
        margin: 1.5rem 0 0.75rem 0;
        line-height: 1.4;
    }}

    .{container_class} h4 {{
        color: #4b5563;
        font-size: 1.2rem;
        font-weight: 600;
        margin: 1.25rem 0 0.5rem 0;
        line-height: 1.4;
    }}
    
    .{container_class} p {{
        font-size: 1.05rem;
        line-height: 1.8;
        margin: 0 0 1.25rem 0;
        color: #374151;
    }}

    .{container_class} ul, .{container_class} ol {{
        margin: 0 0 1.25rem 0;
        padding-left: 1.75rem;
    }}

    .{container_class} li {{
        font-size: 1.05rem;
        line-height: 1.7;
        margin-bottom: 0.75rem;
        color: #374151;
    }}
    
    .{container_class} a {{
        color: {'#1d4ed8' if is_original else '#7c3aed'};
        text-decoration: underline;
        font-weight: 500;
    }}
    
    .{container_class} a:hover {{
        color: {'#1e40af' if is_original else '#6d28d9'};
    }}
    
    .{container_class} strong, .{container_class} b {{
        font-weight: 600;
        color: #1f2937;
    }}
    
    .{container_class} em, .{container_class} i {{
        font-style: italic;
        color: #6b7280;
    }}
    
    .{container_class} blockquote {{
        border-left: 4px solid {'#d1d5db' if is_original else '#7c3aed'};
        margin: 1rem 0;
        padding: 0.75rem 1rem;
        background: {'#f9fafb' if is_original else '#f8fafc'};
        font-style: italic;
        border-radius: 0 4px 4px 0;
    }}
    
    .{container_class} table {{
        width: 100%;
        border-collapse: collapse;
        margin: 1rem 0;
        font-size: 0.9rem;
    }}
    
    .{container_class} th, .{container_class} td {{
        border: 1px solid #d1d5db;
        padding: 0.75rem;
        text-align: left;
    }}
    
    .{container_class} th {{
        background: {'#f3f4f6' if is_original else '#ede9fe'};
        font-weight: 600;
        color: #1f2937;
    }}
    
    .{container_class} code {{
        background: #f3f4f6;
        padding: 0.2rem 0.4rem;
        border-radius: 3px;
        font-family: 'Monaco', 'Menlo', monospace;
        font-size: 0.9rem;
        color: #1f2937;
    }}
    
    /* Improved content subtle styling only (no decorative icons added) */
    .improved-content {{
        box-shadow: 0 4px 6px -1px rgba(124, 58, 237, 0.08);
    }}
    
    /* Responsive design for better width handling */
    @media (max-width: 768px) {{
        .{container_class} {{
            padding: 1rem;
            margin: 0.25rem 0;
            max-height: none;
            width: 100% !important;
            min-width: 100%;
            max-width: 100%;
        }}
        
        .{container_class} h1 {{
            font-size: 1.75rem;
        }}
        
        .{container_class} h2 {{
            font-size: 1.5rem;
        }}
        
        .{container_class} h3 {{
            font-size: 1.25rem;
        }}
    }}
    
    @media (min-width: 1200px) {{
        .{container_class} {{
            padding: 2rem;
            max-height: none;
            width: 100% !important;
            min-width: 100%;
            max-width: 100%;
        }}
    }}
    </style>
    """
    
    # Wrap content in styled container
    formatted_content = f"""
    {css_styles}
    <div class="{container_class}">
        {str(soup)}
    </div>
    """
    
    # Display using Streamlit components with better width handling
    components.html(formatted_content, height=800, scrolling=True)

def display_content_columns_and_suggestions(result: dict, original_text: str):
    with st.container():
        # Use full width with better spacing and responsive design
        col1, col2 = st.columns([1, 1], gap="large")
        
        with col1:
            st.markdown("### Input Content")
            # Create a container with better width control
            with st.container():
                display_formatted_content(original_text, is_original=True)
        
        with col2:
            st.markdown("### Improved Content")
            revised = result.get("revised_content", "")
            # Create a container with better width control
            with st.container():
                display_formatted_content(revised, is_original=False)

            # Download as HTML
            st.download_button(
                label="Download as HTML",
                data=revised,
                file_name="revised_content.html",
                mime="text/html"
            )

            soup = BeautifulSoup(revised, "html.parser")
            doc = docx.Document()

            for elem in soup.recursiveChildGenerator():
                if elem.name:
                    if elem.name == "h1":
                        doc.add_heading(elem.get_text(), level=1)
                    elif elem.name == "h2":
                        doc.add_heading(elem.get_text(), level=2)
                    elif elem.name == "h3":
                        doc.add_heading(elem.get_text(), level=3)
                    elif elem.name == "p":
                        doc.add_paragraph(elem.get_text())
                    elif elem.name == "ul":
                        for li in elem.find_all("li"):
                            doc.add_paragraph(li.get_text(), style="List Bullet")
                    elif elem.name == "ol":
                        for li in elem.find_all("li"):
                            doc.add_paragraph(li.get_text(), style="List Number")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Download as Word document
            st.download_button(
                label="Download as Word",
                data=buffer,
                file_name="revised_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
